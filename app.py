# app.py – OmniRAG: The Multilingual AI Document Assistant (Updated Version)
"""
Features
========
1. Accepts PDF, DOCX, Image, Web URL, YouTube link (auto‑transcript)
2. Generates a Markdown‑formatted summary with adjustable length
3. Opens an ongoing chat powered by Groq (Llama 3) with retrieval over the indexed content
4. Streamlit built-in caching → skips re‑embedding already‑seen docs
5. Exports session report as a formatted DOCX file or a structured JSON file
6. Multi-language support for UI and content (translation via Groq)
7. Text-to-Speech for the summary (with robust Markdown cleaning for natural audio)
8. Indian language support (Telugu, Tamil, Hindi)
9. Enhanced error handling for OCR and document processing
"""

from __future__ import annotations
import os, hashlib, json, tempfile, pathlib, re, io
import streamlit as st
from typing import List, Dict

# ──────────────────────────────────────────────────────────────────────────────
# External libraries
# ──────────────────────────────────────────────────────────────────────────────
from langchain_groq import ChatGroq
from langchain.chains import ConversationalRetrievalChain
from langchain_community.document_loaders import PyPDFLoader, WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings  # Updated import
from langchain.docstore.document import Document
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound
from docx import Document as DocxDoc
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
from gtts import gTTS
import markdown
import pytesseract, requests, bs4

# Explicitly set tesseract path for deployed environments
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

# ──────────────────────────────────────────────────────────────────────────────
# Configuration & Constants
# ──────────────────────────────────────────────────────────────────────────────

# Updated language support - removed German, Japanese, Bengali, French
LANGUAGES = {
    "English": "en", "Español": "es", "हिन्दी": "hi", 
    "తెలుగు": "te", "தமிழ்": "ta"
}

# Updated UI labels for remaining languages
UI_LABELS = {
    "en": {
        "title": "⚡ OmniRAG: Chat with Any Document", 
        "sidebar_title": "1. Add Content", 
        "url_placeholder": "🔗 Public Webpage or YouTube URL", 
        "file_uploader": "📄 Drop Files (PDF, DOCX, PNG, JPG)", 
        "config_header": "2. Configure", 
        "summary_slider": "📝 Summary Length (words)", 
        "process_header": "3. Process", 
        "process_button": "🔄 Process Content", 
        "summary_header": "📝 Summary", 
        "chat_header": "💬 Chat with your content", 
        "chat_placeholder": "Ask anything about the content...", 
        "hear_summary_button": "🔊 Hear Summary", 
        "download_header": "Download Report", 
        "download_json_button": "📥 JSON", 
        "download_docx_button": "📄 DOCX", 
        "initial_info": "👈 Add content in the sidebar and click **Process Content** to begin."
    },
    "es": {
        "title": "⚡ OmniRAG: Chatea con Cualquier Documento", 
        "sidebar_title": "1. Añadir Contenido", 
        "url_placeholder": "🔗 URL de Página Web Pública o YouTube", 
        "file_uploader": "📄 Suelta Archivos (PDF, DOCX, PNG, JPG)", 
        "config_header": "2. Configurar", 
        "summary_slider": "📝 Longitud del Resumen (palabras)", 
        "process_header": "3. Procesar", 
        "process_button": "🔄 Procesar Contenido", 
        "summary_header": "📝 Resumen", 
        "chat_header": "💬 Chatea con tu contenido", 
        "chat_placeholder": "Pregunta lo que sea sobre el contenido...", 
        "hear_summary_button": "🔊 Escuchar Resumen", 
        "download_header": "Descargar Informe", 
        "download_json_button": "📥 JSON", 
        "download_docx_button": "📄 DOCX", 
        "initial_info": "👈 Añade contenido en la barra lateral y haz clic en **Procesar Contenido** para comenzar."
    },
    "hi": {
        "title": "⚡ ओमनीरैग: किसी भी दस्तावेज़ के साथ चैट करें", 
        "sidebar_title": "1. सामग्री जोड़ें", 
        "url_placeholder": "🔗 सार्वजनिक वेबपेज या यूट्यूब यूआरएल", 
        "file_uploader": "📄 फाइलें डालें (पीडीएफ, डॉक्स, पीएनजी, जेपीजी)", 
        "config_header": "2. कॉन्फ़िगर करें", 
        "summary_slider": "📝 सारांश लंबाई (शब्द)", 
        "process_header": "3. संसाधित करें", 
        "process_button": "🔄 सामग्री संसाधित करें", 
        "summary_header": "📝 सारांश", 
        "chat_header": "💬 अपनी सामग्री के साथ चैट करें", 
        "chat_placeholder": "सामग्री के बारे में कुछ भी पूछें...", 
        "hear_summary_button": "🔊 सारांश सुनें", 
        "download_header": "रिपोर्ट डाउनलोड करें", 
        "download_json_button": "📥 JSON", 
        "download_docx_button": "📄 DOCX", 
        "initial_info": "👈 साइडबार में सामग्री जोड़ें और शुरू करने के लिए **सामग्री संसाधित करें** पर क्लिक करें।"
    },
    "te": {
        "title": "⚡ ఓమ్నిరాగ్: ఏదైనా డాక్యుమెంట్‌తో చాట్ చేయండి", 
        "sidebar_title": "1. కంటెంట్ జోడించండి", 
        "url_placeholder": "🔗 పబ్లిక్ వెబ్‌పేజీ లేదా యూట్యూబ్ URL", 
        "file_uploader": "📄 ఫైల్‌లను డ్రాప్ చేయండి (PDF, DOCX, PNG, JPG)", 
        "config_header": "2. కాన్ఫిగర్ చేయండి", 
        "summary_slider": "📝 సారాంశం పొడవు (పదాలు)", 
        "process_header": "3. ప్రాసెస్ చేయండి", 
        "process_button": "🔄 కంటెంట్‌ను ప్రాసెస్ చేయండి", 
        "summary_header": "📝 సారాంశం", 
        "chat_header": "💬 మీ కంటెంట్‌తో చాట్ చేయండి", 
        "chat_placeholder": "కంటెంట్ గురించి ఏదైనా అడగండి...", 
        "hear_summary_button": "🔊 సారాంశం వినండి", 
        "download_header": "నివేదికను డౌన్‌లోడ్ చేయండి", 
        "download_json_button": "📥 JSON", 
        "download_docx_button": "📄 DOCX", 
        "initial_info": "👈 సైడ్‌బార్‌లో కంటెంట్ జోడించి, ప్రారంభించడానికి **కంటెంట్‌ను ప్రాసెస్ చేయండి** క్లిక్ చేయండి."
    },
    "ta": {
        "title": "⚡ ஆம்னிராக்: எந்த ஆவணத்துடனும் அரட்டையடிக்கவும்", 
        "sidebar_title": "1. உள்ளடக்கத்தைச் சேர்க்கவும்", 
        "url_placeholder": "🔗 பொது வலைப்பக்கம் அல்லது YouTube URL", 
        "file_uploader": "📄 கோப்புகளை இழுத்து விடுங்கள் (PDF, DOCX, PNG, JPG)", 
        "config_header": "2. உள்ளமைக்கவும்", 
        "summary_slider": "📝 சுருக்கத்தின் நீளம் (வார்த்தைகள்)", 
        "process_header": "3. செயல்படுத்தவும்", 
        "process_button": "🔄 உள்ளடக்கத்தைச் செயல்படுத்தவும்", 
        "summary_header": "📝 சுருக்கம்", 
        "chat_header": "💬 உங்கள் உள்ளடக்கத்துடன் அரட்டையடிக்கவும்", 
        "chat_placeholder": "உள்ளடக்கம் பற்றி எதையும் கேளுங்கள்...", 
        "hear_summary_button": "🔊 சுருக்கத்தைக் கேட்கவும்", 
        "download_header": "அறிக்கையைப் பதிவிறக்கவும்", 
        "download_json_button": "📥 JSON", 
        "download_docx_button": "📄 DOCX", 
        "initial_info": "👈 பக்க பட்டியில் உள்ளடக்கத்தைச் சேர்த்து, தொடங்க **உள்ளடக்கத்தைச் செயல்படுத்தவும்** என்பதைக் கிளிக் செய்யவும்."
    }
}

# Get Groq API key
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY")
if not GROQ_API_KEY:
    st.error("GROQ_API_KEY not found. Please add it to your .streamlit/secrets.toml file.")
    st.stop()

LLM = ChatGroq(api_key=GROQ_API_KEY, model_name="llama3-8b-8192", temperature=0.2)
EMBEDDER = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")  # Fixed import
SPLITTER = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

# ──────────────────────────────────────────────────────────────────────────────
# Helper Functions
# ──────────────────────────────────────────────────────────────────────────────

def sha256_of_bytes(b: bytes) -> str: 
    return hashlib.sha256(b).hexdigest()

def doc_id_from_source(src: str | bytes) -> str: 
    return sha256_of_bytes(src.encode()) if isinstance(src, str) else sha256_of_bytes(src)

def load_pdf(path: str) -> List[Document]:
    docs = PyPDFLoader(path).load()
    if not docs or all(not doc.page_content.strip() for doc in docs):
        return [Document(page_content="(Empty or scanned PDF — no text found.)", metadata={"source": path})]
    return docs

def load_docx(path: str) -> List[Document]: 
    return [Document(page_content="\n".join(p.text for p in DocxDoc(path).paragraphs), metadata={"source": path})]

def load_image(path: str) -> List[Document]:
    try:
        text = pytesseract.image_to_string(Image.open(path))
        if not text.strip():
            raise ValueError("OCR found no readable text in the image.")
        return [Document(page_content=text, metadata={"source": path})]
    except Exception as e:
        return [Document(page_content=f"(Error reading image: {e})", metadata={"source": path})]

def load_youtube(url: str) -> List[Document]:
    try:
        vid_id = re.findall(r"(?<=v=|be/)[^&#?]+", url)[0]
        transcript = YouTubeTranscriptApi.get_transcript(vid_id)
        text = "\n".join([t['text'] for t in transcript])
    except NoTranscriptFound: 
        text = "(No transcript available for this video.)"
    except Exception as e: 
        text = f"(Error fetching transcript: {e})"
    return [Document(page_content=text, metadata={"source": url})]

def load_web(url: str) -> List[Document]:
    try:
        loader = WebBaseLoader(url)
        loader.requests_per_second = 1
        docs = loader.load()
        for doc in docs:
            soup = bs4.BeautifulSoup(doc.page_content, "html.parser")
            [s.extract() for s in soup(['script', 'style', 'nav', 'footer', 'aside'])]
            doc.page_content = soup.get_text(separator="\n", strip=True)
        return docs
    except Exception as e: 
        return [Document(page_content=f"Error loading web page: {e}", metadata={"source": url})]

# Updated caching function using Streamlit's built-in caching
@st.cache_data
def get_cache_key(cache_key_components: List[str]) -> str:
    return sha256_of_bytes("||".join(sorted(cache_key_components)).encode())

@st.cache_resource
def get_or_build_index(_docs: List[Document], doc_key: str) -> FAISS:
    """Build FAISS index with Streamlit caching"""
    with st.spinner("⚙️ Indexing content... (this may take a moment)"):
        chunks = SPLITTER.split_documents(_docs)
        if not chunks:
            raise ValueError("Indexing failed: No text chunks were created from the document.")
        index = FAISS.from_documents(chunks, EMBEDDER)
    return index

@st.cache_data
def get_cached_summary(content_hash: str, summary_length: int, language: str) -> str:
    """Cache summaries using Streamlit's built-in caching"""
    return None  # This will be populated when summaries are generated

def translate_text(text: str, target_lang_name: str, llm: ChatGroq) -> str:
    if target_lang_name == "English": 
        return text
    prompt = f"Translate the following text accurately to {target_lang_name}. Provide only the translated text, without any additional commentary or explanations.\n\nText to translate:\n---\n{text}"
    return llm.invoke(prompt).content

def clean_markdown_for_tts(markdown_text: str) -> str:
    html = markdown.markdown(markdown_text)
    soup = bs4.BeautifulSoup(html, "html.parser")
    plain_text = soup.get_text(separator=" ")
    return ' '.join(plain_text.split())

def text_to_speech_bytes(text: str, lang_code: str) -> io.BytesIO:
    tts = gTTS(text=text, lang=lang_code, slow=False)
    fp = io.BytesIO()
    tts.write_to_fp(fp)
    fp.seek(0)
    return fp

# Updated DOCX report function (replaces PDF)
def create_docx_report(summary: str, chat_history: List[Dict], labels: Dict) -> bytes:
    """Create a DOCX report with summary and chat history"""
    doc = DocxDoc()
    
    # Set document title
    title = doc.add_heading('OmniRAG Session Report', 0)
    title.alignment = 1  # Center alignment
    
    # Add summary section
    doc.add_heading(labels['summary_header'], level=1)
    
    # Clean markdown and add summary paragraphs
    summary_text = clean_markdown_for_tts(summary)
    summary_para = doc.add_paragraph(summary_text)
    summary_para.style = 'Normal'
    
    # Add spacing
    doc.add_paragraph("")
    
    # Add chat history section
    doc.add_heading(labels['chat_header'], level=1)
    
    for turn in chat_history:
        role = turn["role"].capitalize()
        content = clean_markdown_for_tts(turn["display_content"])
        
        # Add role as bold paragraph
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"{role}:")
        role_run.bold = True
        
        # Add content
        content_para = doc.add_paragraph(content)
        content_para.style = 'Normal'
        
        # Add spacing between turns
        doc.add_paragraph("")
    
    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

# ──────────────────────────────────────────────────────────────────────────────
# Streamlit UI
# ──────────────────────────────────────────────────────────────────────────────

st.set_page_config(page_title="OmniRAG: Multilingual AI Assistant", layout="wide")

# Initialize session state
if "chat_history" not in st.session_state: 
    st.session_state.chat_history = []
if "current_key" not in st.session_state: 
    st.session_state.current_key = None
if "language" not in st.session_state: 
    st.session_state.language = "English"

# Language selection in sidebar
selected_lang_name = st.sidebar.selectbox(
    "Language", 
    options=list(LANGUAGES.keys()), 
    index=list(LANGUAGES.keys()).index(st.session_state.language)
)
if selected_lang_name != st.session_state.language:
    st.session_state.language = selected_lang_name
    st.rerun()

# Get language code and UI labels
lang_code = LANGUAGES[st.session_state.language]
labels = UI_LABELS.get(lang_code, UI_LABELS["en"])

st.title(labels["title"])
st.markdown("<sub>Powered by Groq LPU™ Inference Engine for blazing-fast, multilingual responses</sub>", unsafe_allow_html=True)

with st.sidebar:
    st.header(labels["sidebar_title"])
    url_input = st.text_input(labels["url_placeholder"])
    files = st.file_uploader(labels["file_uploader"], type=["pdf", "docx", "png", "jpg", "jpeg"], accept_multiple_files=True)
    
    st.header(labels["config_header"])
    summary_length = st.slider(labels["summary_slider"], 100, 1000, 300, 50)
    
    st.header(labels["process_header"])
    process_btn = st.button(labels["process_button"], type="primary")

if process_btn:
    if not url_input and not files:
        st.error("Please provide a URL or upload at least one file.")
        st.stop()

    all_docs, cache_key_components = [], []
    try:
        if url_input:
            with st.spinner(f"Fetching content from {url_input}..."):
                is_youtube = "youtube.com/watch?v=" in url_input or "youtu.be/" in url_input
                docs = load_youtube(url_input) if is_youtube else load_web(url_input)
                all_docs.extend(docs)
                cache_key_components.append(url_input)
                
        for f in files or []:
            with st.spinner(f"Processing file: {f.name}..."):
                file_bytes = f.read()
                with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{f.name}") as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                
                if f.type == "application/pdf": 
                    docs = load_pdf(tmp_path)
                elif f.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document": 
                    docs = load_docx(tmp_path)
                elif f.type.startswith("image/"): 
                    docs = load_image(tmp_path)
                else: 
                    st.warning(f"Unsupported file type: {f.name} ({f.type})")
                    docs = []
                
                os.unlink(tmp_path)
                all_docs.extend(docs)
                cache_key_components.append(doc_id_from_source(file_bytes))
                
    except Exception as e:
        st.error(f"An error occurred during processing: {e}")
        st.stop()

    if not all_docs: 
        st.error("No valid content could be processed.")
        st.stop()

    # Use Streamlit caching for the key generation and index building with error handling
    full_key = get_cache_key(cache_key_components)
    try:
        index = get_or_build_index(all_docs, full_key)
    except Exception as e:
        st.error(f"Indexing failed: {e}")
        st.stop()

    # Generate summary with caching
    @st.cache_data
    def generate_summary(content_key: str, length: int) -> str:
        chain = ConversationalRetrievalChain.from_llm(LLM, index.as_retriever())
        prompt = f"Summarise the provided content in clear Markdown. Use headings (##), bullet points, and **bold** for key terms. The summary should be approximately {length} words."
        return chain.invoke({"question": prompt, "chat_history": []})["answer"]

    # Generate English summary
    original_summary = generate_summary(full_key, summary_length)

    # Translate summary if needed
    if st.session_state.language != "English":
        @st.cache_data
        def get_translated_summary(summary_text: str, target_language: str) -> str:
            return translate_text(summary_text, target_language, LLM)
        
        display_summary = get_translated_summary(original_summary, st.session_state.language)
    else:
        display_summary = original_summary

    st.session_state.summary = display_summary
    st.session_state.chat_chain = ConversationalRetrievalChain.from_llm(LLM, index.as_retriever(search_kwargs={"k": 4}))
    st.session_state.chat_history = []
    st.session_state.current_key = full_key

# Display content if it has been processed
if st.session_state.current_key:
    st.markdown(f"## {labels['summary_header']}")
    
    # Check if summary is blank and show warning
    if not st.session_state.summary.strip():
        st.warning("Summary generation returned no meaningful content.")
    else:
        st.markdown(st.session_state.summary, unsafe_allow_html=True)
    
    col1, col2 = st.columns([0.8, 0.2])
    with col2:
        if st.button(labels['hear_summary_button'], use_container_width=True):
            with st.spinner("Generating audio..."):
                plain_text_summary = clean_markdown_for_tts(st.session_state.summary)
                audio_bytes = text_to_speech_bytes(plain_text_summary, lang_code)
                st.audio(audio_bytes, format='audio/mp3')

    st.divider()
    st.markdown(f"### {labels['chat_header']}")

    for turn in st.session_state.chat_history:
        st.chat_message(turn["role"]).markdown(turn["display_content"], unsafe_allow_html=True)

    user_q = st.chat_input(labels["chat_placeholder"])
    if user_q:
        st.chat_message("user").markdown(user_q)
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                model_history = [(turn["role"], turn["original_content"]) for turn in st.session_state.chat_history]
                response = st.session_state.chat_chain.invoke({"question": user_q, "chat_history": model_history})
                original_answer = response["answer"]
                
                display_answer = translate_text(original_answer, st.session_state.language, LLM) if st.session_state.language != "English" else original_answer
                st.markdown(display_answer, unsafe_allow_html=True)
        
        st.session_state.chat_history.append({"role": "user", "original_content": user_q, "display_content": user_q})
        st.session_state.chat_history.append({"role": "assistant", "original_content": original_answer, "display_content": display_answer})

    if st.session_state.summary:
        with st.sidebar:
            st.divider()
            st.header(labels["download_header"])
            
            report_data_json = {
                "summary": st.session_state.summary,
                "chat_history": [turn["display_content"] for turn in st.session_state.chat_history]
            }
            
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label=labels["download_json_button"],
                    data=json.dumps(report_data_json, indent=2, ensure_ascii=False),
                    file_name=f"report_{st.session_state.current_key[:8]}_{lang_code}.json",
                    mime="application/json",
                    use_container_width=True
                )
            
            with col2:
                docx_data = create_docx_report(st.session_state.summary, st.session_state.chat_history, labels)
                st.download_button(
                    label=labels["download_docx_button"],
                    data=docx_data,
                    file_name=f"report_{st.session_state.current_key[:8]}_{lang_code}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

else:
    st.info(labels["initial_info"])
